from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

from collections import OrderedDict
from pathlib import Path
import calendar
import datetime as dt
import os



HEADER_FILL = '5b95f9'
COLOR1 = 'B6CFFC'

class WordTable():

	def __init__(self, duty_dates, total_days, filename):
		self.duty_dates = OrderedDict(zip(list(map(self.convert_dates_to_text, duty_dates.keys())), duty_dates.values()))
		self.total_days = total_days
		self.path = Path('backend/word_files/') / (filename + ".docx")

	def convert_dates_to_text(self, date_string):
		shift_start, shift_end = list(map(lambda x : dt.datetime.strptime(x, '%Y-%m-%d'), date_string.split(' - ')))
		shift_start_text = "{} {}".format(calendar.month_name[shift_start.month], shift_start.day)
		shift_end_text = "{} {}".format(calendar.month_name[shift_end.month], shift_end.day)
		text = "{} - {}".format(shift_start_text, shift_end_text)
		
		return text
		
        
	def write_to_table(self):
		document = Document()
		
		drows = len(self.duty_dates) + 1
		trows = len(self.total_days) + 1
		
		styles = document.styles
		
		duty_table = document.add_table(rows=drows, cols=5)
		document.add_page_break()
		total_table = document.add_table(rows=trows, cols=3)

		duty_table.style = styles['Table Grid']
		total_table.style = styles['Table Grid']

		dheaders = ['Week', 'Dates', 'RD', 'Management', 'Central Office']
		theaders = ['Name', 'Weekdays', 'Weekends']

		#-----Table Formatting-----
		for i in range(1, drows, 2):
			row = duty_table.rows[i].cells[3]
			try:
				row2 = duty_table.rows[i + 1].cells[3]
				row.merge(row2)
			except IndexError:
				pass

		duty_table.cell(0, 1).width = Inches(2.5)
		duty_table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER
		

		total_table.cell(0, 1).width = Inches(2.5)
		total_table.cell(0, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

		for cell in range(5):
			header_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), HEADER_FILL))
			try:
				duty_table.rows[0].cells[cell]._tc.get_or_add_tcPr().append(header_fill)
				total_table.rows[0].cells[cell]._tc.get_or_add_tcPr().append(header_fill)
			
			except IndexError:
				pass

		for i in range(1, drows, 2):
			row = duty_table.rows[i]
			
			for j in range(5):
				cell_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), COLOR1))
				row.cells[j]._tc.get_or_add_tcPr().append(cell_fill)

		for i in range(1, trows, 2):
			row = total_table.rows[i]

			for j in range(3):
				cell_fill = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), COLOR1))
				row.cells[j]._tc.get_or_add_tcPr().append(cell_fill)

		#-----Fill Header Info-----
		for i in range(len(dheaders)):
			text = dheaders[i]
			p = duty_table.rows[0].cells[i].add_paragraph(text + "\n")
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER
		
		for i in range(len(theaders)):
			text = theaders[i]
			p = total_table.rows[0].cells[i].add_paragraph(text + "\n")
			p.alignment = WD_ALIGN_PARAGRAPH.CENTER	

		#------Write to Table-------
		for i in range(len(self.duty_dates.items())):
			date_info = list(self.duty_dates.items())[i]
			row_info = ["Week %s" % (i + 1), date_info[0], date_info[1]]

			for j in range(len(row_info)):
				duty_table.rows[i + 1].cells[j].text = row_info[j]

		for i in range(len(self.total_days)):
			x = list(self.total_days.items())
			row_info = [x[i][0], x[i][1]['weekdays'], x[i][1]['weekends']]
			
			for j in range(len(row_info)):
				total_table.rows[i + 1].cells[j].text = str(row_info[j])

		document.save(self.path)